"""
Build Chapters 4–5 Word document with tables, captions, and embedded figures.
Placeholders remain where the user must paste a screenshot (e.g. dashboard).
"""
from __future__ import annotations

import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor, Cm

BASE = os.path.dirname(os.path.abspath(__file__))
FIG = os.path.join(BASE, "figures")
OUT = os.path.join(BASE, "report", "Chapters_4_and_5_System_Implementation_Results_and_Conclusion.docx")


def set_run_font(run, name="Times New Roman", size=12, bold=False, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_para(doc, text, *, style=None, bold=False, italic=False, size=12, space_after=8, space_before=0, align=None):
    p = doc.add_paragraph()
    if style:
        p.style = style
    if align:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)
    return p


def add_heading_custom(doc, text, level=1):
    # Use built-in heading styles for TOC
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        set_run_font(run, size=14 if level == 1 else 12, bold=True)
    h.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    h.paragraph_format.space_after = Pt(8)
    return h


def add_caption(doc, text):
    p = add_para(doc, text, italic=True, size=11, space_after=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    return p


def add_placeholder(doc, label):
    p = add_para(
        doc,
        f"[INSERT IMAGE HERE: {label}]",
        bold=True,
        italic=True,
        size=11,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=12,
        space_after=12,
    )
    # light border via shading simulation - just leave clear instruction
    return p


def try_add_picture(doc, filename, width_in=5.8, caption=None):
    path = os.path.join(FIG, filename)
    if os.path.isfile(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=Inches(width_in))
        if caption:
            add_caption(doc, caption)
        return True
    add_placeholder(doc, filename)
    if caption:
        add_caption(doc, caption)
    return False


def set_cell_shading(cell, fill_hex):
    tc = cell._tePr if hasattr(cell, "_tePr") else cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(h)
        set_run_font(run, size=10, bold=True)
        set_cell_shading(hdr_cells[i], "D9E2F3")
    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            cell = table.rows[r_i + 1].cells[c_i]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            set_run_font(run, size=10)
    if col_widths:
        for row in table.rows:
            for idx, w in enumerate(col_widths):
                row.cells[idx].width = Inches(w)
    doc.add_paragraph()
    return table


def body(doc, text):
    return add_para(doc, text, size=12, space_after=10)


def build():
    doc = Document()

    # Page setup roughly A4-friendly margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    # Title block
    add_para(
        doc,
        "Explainable Artificial Intelligence (XAI) Framework for Enhancing Managerial Decision-Making in SMEs",
        bold=True,
        size=14,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=6,
    )
    add_para(
        doc,
        "Chapters 4 and 5 — System Implementation, Results, Conclusion and Recommendations",
        bold=True,
        size=12,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=6,
    )
    add_para(
        doc,
        "Paste these chapters into the main dissertation after Chapter 3. Figure files live in the project figures/ folder.",
        italic=True,
        size=10,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=18,
    )

    # ===================== CHAPTER 4 =====================
    add_heading_custom(doc, "CHAPTER FOUR: SYSTEM IMPLEMENTATION AND RESULTS", level=1)

    add_heading_custom(doc, "4.1 Introduction", level=2)
    body(
        doc,
        "This chapter presents the implementation and empirical evaluation of the proposed Explainable Artificial Intelligence (XAI) framework for enhancing managerial decision-making in small and medium enterprises (SMEs). Whereas Chapter Three specified the research design, data source, machine learning models, explainability techniques and evaluation metrics, the present chapter reports what was actually built, measured and interpreted.",
    )
    body(
        doc,
        "The implementation follows the experimental pipeline developed in the project notebook (XAI_Bus.ipynb) and operationalised through saved model artefacts, evaluation reports, visualisations, a FastAPI inference service and a managerial decision-support dashboard. The chapter mirrors that pipeline: implementation environment; dataset and preprocessing; exploratory analysis; model development and evaluation; feature importance and error analysis; SHAP and LIME explainability; the decision-support framework; and discussion of findings against the research objectives and questions.",
    )
    body(
        doc,
        "All quantitative results below are taken from the held-out test evaluation, cross-validation outputs and explainability artefacts generated during implementation.",
    )

    add_heading_custom(doc, "4.2 Implementation Environment", level=2)
    body(
        doc,
        "The framework was implemented in Python within an interactive Jupyter / Google Colab environment to support rapid experimentation, visualisation and artefact export. Core libraries included Pandas and NumPy for data handling; Scikit-learn for preprocessing, Decision Tree, Random Forest, metrics and cross-validation; imbalanced-learn (SMOTE) for training-set class rebalancing; XGBoost for gradient boosting classification; SHAP and LIME for explainability; Matplotlib for figures; and Joblib/Pickle for artefact persistence. A FastAPI backend and Next.js dashboard operationalise inference and managerial reporting.",
    )
    body(
        doc,
        "Reproducibility was supported by fixing random_state=42 for the stratified split, SMOTE and model initialisation, and by exporting artefacts to structured directories (figures/, metrics/, reports/, and serialised .pkl files).",
    )
    add_para(doc, "Table 4.1: Summary of model and pipeline hyperparameters", bold=True, size=11, space_after=6)
    add_table(
        doc,
        ["Component", "Setting"],
        [
            ["Train–test split", "80% / 20%, stratified, random_state=42"],
            ["SMOTE", "Training data only, random_state=42"],
            ["Decision Tree", "Default scikit-learn, random_state=42"],
            ["Random Forest", "n_estimators=100, random_state=42"],
            ["XGBoost", "n_estimators=100, learning_rate=0.1, eval_metric=logloss, random_state=42"],
            ["Cross-validation", "5-fold StratifiedKFold, scoring=F1"],
            ["LIME", "Tabular explainer, top 5 local features per case"],
            ["SHAP", "TreeExplainer on best model"],
        ],
        col_widths=[2.2, 4.3],
    )

    add_heading_custom(doc, "4.3 Dataset Description", level=2)
    body(
        doc,
        "The empirical study uses the IBM HR Analytics Employee Attrition and Performance dataset, a publicly available Kaggle benchmark widely used for workforce analytics and binary classification. The dataset is synthetic but designed to reflect realistic organisational attributes, supporting methodological demonstration without exposing real employee identities.",
    )
    add_para(doc, "Table 4.2: Dataset characteristics", bold=True, size=11, space_after=6)
    add_table(
        doc,
        ["Characteristic", "Description"],
        [
            ["Number of observations", "1,470 employees"],
            ["Number of variables (raw)", "35 features"],
            ["Format", "CSV"],
            ["Target variable", "Attrition (Yes/No → 1/0)"],
            ["Missing values", "None reported in source file"],
            ["Learning problem", "Binary classification"],
            ["Source", "Kaggle – IBM HR Analytics Employee Attrition"],
        ],
        col_widths=[2.5, 4.0],
    )
    body(
        doc,
        "Predictor variables span demographics, job and organisational factors, compensation, satisfaction/engagement and tenure. Although the case domain is employee attrition, the XAI framework (predict → explain → recommend) is designed to generalise to other SME tabular decision problems.",
    )

    add_heading_custom(doc, "4.4 Data Preprocessing", level=2)
    body(
        doc,
        "Preprocessing prepared the raw dataset for supervised learning while preserving feature interpretability for SHAP and LIME. Non-informative fields EmployeeNumber and Over18 were dropped. The target Attrition was mapped Yes→1, No→0. A ColumnTransformer applied StandardScaler to numerical features and OneHotEncoder (drop='first') to categorical features, producing 46 model features. Data were split 80/20 with stratification. Because attrition is a minority class, SMOTE was applied only to the training data after preprocessing, avoiding leakage into the test set. Artefacts saved for inference included preprocessor.pkl and feature_names.pkl. The held-out test set contained n = 294 employees.",
    )
    add_para(doc, "Table 4.3: Preprocessing outcomes", bold=True, size=11, space_after=6)
    add_table(
        doc,
        ["Stage", "Outcome"],
        [
            ["Raw shape", "1,470 × 35 (approx.)"],
            ["After drop of ID/constants", "Reduced feature set; target encoded"],
            ["After one-hot + scaling", "46 model features"],
            ["Split", "Stratified 80/20"],
            ["SMOTE", "Training rebalanced; test left natural"],
        ],
        col_widths=[2.5, 4.0],
    )

    add_heading_custom(doc, "4.5 Exploratory Data Analysis", level=2)
    body(
        doc,
        "Exploratory analysis confirmed data quality and motivated modelling choices. Attrition is imbalanced: most employees stay and a minority leave (approximately 16.1% attrition). This imbalance justifies multi-metric evaluation (precision, recall, F1, ROC-AUC) rather than accuracy alone, and motivates SMOTE on the training set only. No missing values or duplicate rows were observed in the cleaned working frame.",
    )
    try_add_picture(
        doc,
        "attrition_distribution.png",
        width_in=5.5,
        caption="Figure 4.2: Distribution of employee attrition in the IBM HR dataset, showing class imbalance between stayers and leavers.",
    )
    body(
        doc,
        "A correlation matrix of numerical predictors was examined to understand linear associations among tenure, income and job-level related variables. Tree ensembles handle such correlations through interactive splits without requiring feature independence assumptions.",
    )
    try_add_picture(
        doc,
        "correlation_matrix.png",
        width_in=5.8,
        caption="Figure 4.3: Correlation matrix of numerical predictors (computed on test-set feature values used in error analysis).",
    )

    add_heading_custom(doc, "4.6 Model Development", level=2)
    body(
        doc,
        "Three supervised classifiers were trained on the SMOTE-resampled training matrix and compared on the untouched stratified test set: (1) Decision Tree as a transparent baseline; (2) Random Forest as a bagged ensemble; and (3) XGBoost as gradient-boosted trees. All models used a fixed random seed. After training, class labels and attrition probabilities were generated for every test employee. The model with the highest ROC-AUC on the test set—XGBoost—was selected for explainability and decision support and serialised as xgb_model.pkl.",
    )

    add_heading_custom(doc, "4.7 Model Evaluation", level=2)
    add_heading_custom(doc, "4.7.1 Hold-out test performance", level=3)
    add_para(doc, "Table 4.4: Model performance comparison on the held-out test set (n = 294)", bold=True, size=11, space_after=6)
    add_table(
        doc,
        ["Model", "Accuracy", "Precision", "Recall", "F1-Score", "ROC-AUC"],
        [
            ["Decision Tree", "0.7857", "0.3261", "0.3191", "0.3226", "0.5968"],
            ["Random Forest", "0.8401", "0.5000", "0.2128", "0.2985", "0.7805"],
            ["XGBoost", "0.8776", "0.7200", "0.3830", "0.5000", "0.8041"],
        ],
    )
    body(
        doc,
        "XGBoost achieved the best overall discrimination (ROC-AUC = 0.8041) and the highest precision (0.72) and F1 (0.50) among the three models. Accuracy alone would overstate Decision Tree quality under imbalance. Random Forest improved accuracy and AUC relative to the single tree but produced low recall (0.21). XGBoost still shows modest recall (0.38), which is material for retention: false negatives (missed leavers) carry high managerial cost.",
    )

    add_heading_custom(doc, "4.7.2 Cross-validation", level=3)
    add_para(doc, "Table 4.5: Five-fold stratified cross-validation (F1 on SMOTE-resampled training data)", bold=True, size=11, space_after=6)
    add_table(
        doc,
        ["Model", "Mean F1", "Std Dev"],
        [
            ["Decision Tree", "0.8549", "0.0174"],
            ["Random Forest", "0.9323", "0.0112"],
            ["XGBoost", "0.9270", "0.0073"],
        ],
    )
    body(
        doc,
        "Cross-validation F1 is high for the ensembles. Random Forest slightly leads on mean F1 while XGBoost shows the lowest variance. Differences between CV F1 and test F1 are expected because CV was scored on resampled training folds, whereas Table 4.4 reflects the natural imbalanced test distribution. Final selection prioritised test ROC-AUC and precision–recall behaviour relevant to deployment.",
    )

    add_heading_custom(doc, "4.7.3 ROC curves and confusion matrix", level=3)
    try_add_picture(
        doc,
        "roc_curve_comparison.png",
        width_in=5.5,
        caption="Figure 4.4: ROC curves for Decision Tree, Random Forest and XGBoost on the held-out test set.",
    )
    try_add_picture(
        doc,
        "confusion_matrix.png",
        width_in=4.8,
        caption="Figure 4.5: Confusion matrix for the selected XGBoost model on the test set.",
    )
    add_para(doc, "Table 4.6: Confusion counts for XGBoost (default decision threshold)", bold=True, size=11, space_after=6)
    add_table(
        doc,
        ["Outcome", "Count", "Meaning"],
        [
            ["True Negative (TN)", "240", "Correctly predicted stayers"],
            ["True Positive (TP)", "18", "Correctly predicted leavers"],
            ["False Negative (FN)", "29", "Leavers predicted as stayers"],
            ["False Positive (FP)", "7", "Stayers predicted as leavers"],
            ["Total", "294", "Full test set"],
        ],
    )

    add_heading_custom(doc, "4.7.4 Precision–recall and threshold sensitivity", level=3)
    body(
        doc,
        "Default classification uses probability threshold 0.5. Because attrition is rare, the implementation also examined the precision–recall trade-off and an F1-maximising threshold. Threshold tuning allows SME policy to favour higher recall (catch more potential leavers) or higher precision (fewer false alarms). The decision-support interface can expose risk bands (high / moderate / low) rather than a single hard label.",
    )
    try_add_picture(
        doc,
        "threshold_optimization.png",
        width_in=5.5,
        caption="Figure 4.6: Precision and recall as functions of the classification threshold, with the F1-oriented operating point highlighted.",
    )

    add_heading_custom(doc, "4.8 Feature Importance", level=2)
    body(
        doc,
        "Built-in tree importance provides a first global view of which encoded features the models rely on. Figure 4.7 compares top features across Decision Tree, Random Forest and XGBoost. For the selected XGBoost model, gain-based importance ranked Overtime highest, followed by job level, stock option level, frequent travel and marital status encodings.",
    )
    try_add_picture(
        doc,
        "feature_importance.png",
        width_in=5.5,
        caption="Figure 4.7: Top-15 feature importances for Decision Tree, Random Forest and XGBoost.",
    )
    try_add_picture(
        doc,
        "feature_importance_gain_bar.png",
        width_in=5.5,
        caption="Figure 4.7b: XGBoost gain-based feature importance ranking (Top 12).",
    )
    add_para(doc, "Table 4.7: XGBoost gain-based feature importance ranking (Top 10)", bold=True, size=11, space_after=6)
    add_table(
        doc,
        ["Rank", "Feature (pipeline name)", "Manager-readable label", "Gain"],
        [
            ["1", "cat__OverTime_Yes", "Works overtime", "31.68"],
            ["2", "num__JobLevel", "Job level", "15.24"],
            ["3", "num__StockOptionLevel", "Stock option level", "11.74"],
            ["4", "cat__BusinessTravel_Travel_Frequently", "Travels frequently", "9.48"],
            ["5", "cat__MaritalStatus_Single", "Single", "8.11"],
            ["6", "cat__EducationField_Medical", "Education field: Medical", "7.70"],
            ["7", "num__TotalWorkingYears", "Total working years", "7.68"],
            ["8", "cat__JobRole_Sales Executive", "Sales Executive role", "6.76"],
            ["9", "cat__JobRole_Laboratory Technician", "Laboratory Technician role", "6.63"],
            ["10", "cat__JobRole_Research Scientist", "Research Scientist role", "6.37"],
        ],
    )
    body(
        doc,
        "Split-weight importance instead emphasised continuous variables such as MonthlyIncome, Age, DailyRate and DistanceFromHome. Frequency of splits is not identical to average gain. The dissertation therefore reports both model-centric importance and SHAP attributions (Section 4.10), which are better suited to managerial communication of individual predictions.",
    )

    add_heading_custom(doc, "4.9 Error Analysis", level=2)
    body(
        doc,
        "Error analysis examines who is misclassified and why that matters for SME managers. On the test set, XGBoost produced 240 true negatives, 18 true positives, 29 false negatives and 7 false positives.",
    )
    try_add_picture(
        doc,
        "error_analysis.png",
        width_in=5.0,
        caption="Figure 4.8: Counts of true positives, true negatives, false positives and false negatives for XGBoost on the test set.",
    )
    add_para(doc, "Table 4.9: Outcome summary (XGBoost test set)", bold=True, size=11, space_after=6)
    add_table(
        doc,
        ["Prediction outcome", "Count", "Share of test set"],
        [
            ["True Negative", "240", "81.6%"],
            ["False Negative", "29", "9.9%"],
            ["True Positive", "18", "6.1%"],
            ["False Positive", "7", "2.4%"],
        ],
    )
    body(
        doc,
        "True positives (correctly flagged leavers) were strongly associated with OverTime = Yes (16 of 18), relatively lower mean monthly income (~3,773) and younger mean age (~32), with many single employees. False negatives mixed overtime and non-overtime cases, with higher mean income (~5,045) and age (~38)—employees who may leave for subtler reasons not fully captured by the strongest global drivers. False positives were few (7). These patterns motivate threshold policies, human-in-the-loop review of borderline probabilities, and local XAI before high-stakes retention actions.",
    )

    add_heading_custom(doc, "4.10 SHAP Analysis", level=2)
    body(
        doc,
        "SHAP (SHapley Additive exPlanations) attributes the model’s prediction to features using cooperative game theory (Lundberg and Lee, 2017). A TreeExplainer was fitted to the trained XGBoost model. Global plots summarise average feature impact; local waterfall plots explain individual employees.",
    )
    try_add_picture(
        doc,
        "shap_summary_plot.png",
        width_in=5.8,
        caption="Figure 4.9: SHAP summary (beeswarm) plot for the XGBoost attrition model.",
    )
    try_add_picture(
        doc,
        "shap_bar_plot.png",
        width_in=5.5,
        caption="Figure 4.10: Global SHAP feature importance (mean absolute SHAP value) for the XGBoost attrition model.",
    )
    add_para(doc, "Table 4.10: Top global attrition drivers (SHAP)", bold=True, size=11, space_after=6)
    add_table(
        doc,
        ["Rank", "Feature (pipeline name)", "Manager-readable label", "Mean |SHAP|"],
        [
            ["1", "cat__OverTime_Yes", "Works overtime", "0.9869"],
            ["2", "num__StockOptionLevel", "Stock option level", "0.5120"],
            ["3", "cat__BusinessTravel_Travel_Frequently", "Travels frequently", "0.3364"],
            ["4", "num__NumCompaniesWorked", "Number of companies worked", "0.3147"],
            ["5", "num__JobSatisfaction", "Job satisfaction", "0.2956"],
            ["6", "num__MonthlyIncome", "Monthly income", "0.2924"],
            ["7", "num__Age", "Age", "0.2884"],
            ["8", "num__YearsWithCurrManager", "Years with current manager", "0.2849"],
            ["9", "num__DistanceFromHome", "Distance from home", "0.2638"],
            ["10", "num__EnvironmentSatisfaction", "Environment satisfaction", "0.2442"],
        ],
    )
    body(
        doc,
        "Overtime is by far the strongest global driver (mean |SHAP| 0.987, roughly double the next feature), followed by a cluster of compensation, tenure and satisfaction attributes. The SHAP ranking reconciles the gain-based view (Table 4.7) with split-weight importance: categorical risk switches (overtime, travel, stock options) and continuous HR variables (income, age, tenure) both matter, but attribution-based SHAP puts overtime clearly first, which is directly communicable to managers.",
    )
    try_add_picture(
        doc,
        "shap_dependence_plot.png",
        width_in=5.5,
        caption="Figure 4.11: SHAP dependence plot for Monthly Income versus attributed attrition risk.",
    )
    try_add_picture(
        doc,
        "shap_waterfall_sample.png",
        width_in=5.5,
        caption="Figure 4.12: SHAP waterfall plot for a single test employee, decomposing the prediction from base value through feature contributions.",
    )
    body(
        doc,
        "Dependence and waterfall plots help managers see non-linear effects and case-level logic rather than a single coefficient, operationalising “glass box” communication for organisational decisions.",
    )

    add_heading_custom(doc, "4.11 LIME Analysis", level=2)
    body(
        doc,
        "LIME explains individual predictions by fitting a simple local surrogate around the instance (Ribeiro, Singh and Guestrin, 2016). A LimeTabularExplainer was trained on the SMOTE-resampled training matrix. Three test cases were explained with the top five local features each.",
    )
    try_add_picture(doc, "lime_case_1.png", width_in=5.2, caption="Figure 4.13: LIME local explanation — Case 1.")
    try_add_picture(doc, "lime_case_2.png", width_in=5.2, caption="Figure 4.14: LIME local explanation — Case 2.")
    try_add_picture(doc, "lime_case_3.png", width_in=5.2, caption="Figure 4.15: LIME local explanation — Case 3.")
    add_para(doc, "Table 4.11: LIME case comparison (top five local features per case)", bold=True, size=11, space_after=6)
    add_table(
        doc,
        ["Case", "Predicted class", "Top local factors (LIME weights)", "Managerial reading"],
        [
            [
                "1",
                "Stay",
                "No overtime (−0.127); no frequent travel (−0.067); female (−0.034); R&D department (+0.033); not a Laboratory Technician (−0.027)",
                "Low-risk profile: absence of the two strongest risk switches (overtime, frequent travel) dominates; routine monitoring only",
            ],
            [
                "2",
                "Stay",
                "Overtime (+0.120); no frequent travel (−0.067); higher monthly income (−0.055); no stock options (+0.030); not a Laboratory Technician (−0.029)",
                "Borderline monitor case: overtime and missing stock options raise risk, but income and travel profile pull the other way",
            ],
            [
                "3",
                "Stay",
                "No overtime (−0.124); no frequent travel (−0.067); no stock options (+0.031); female (−0.030); not a Laboratory Technician (−0.030)",
                "Low-risk profile similar to Case 1; stock-option gap is the only material upward driver",
            ],
        ],
    )
    body(
        doc,
        "In all three sampled cases the local explanations are dominated by the same two factors that lead the global rankings—overtime and frequent business travel—illustrating that the local and global views are mutually consistent. Case 2 shows how LIME surfaces a genuinely mixed profile (overtime risk offset by income and travel factors), where a manager should review the employee before acting.",
    )
    add_para(doc, "Table 4.13: SHAP–LIME consistency audit for a sampled test employee", bold=True, size=11, space_after=6)
    add_table(
        doc,
        ["Feature", "SHAP value", "LIME weight"],
        [
            ["Overtime", "−1.0895", "−0.1250"],
            ["Monthly income", "+0.8777", "+0.0235"],
            ["Years with current manager", "+0.8657", "+0.0019"],
            ["Stock option level", "−0.7406", "+0.0269"],
            ["Total working years", "+0.6696", "−0.0013"],
            ["Frequent business travel", "−0.4299", "−0.0713"],
            ["Age", "+0.3584", "+0.0143"],
            ["Job role: Sales Representative", "+0.3013", "+0.0016"],
            ["Environment satisfaction", "−0.2599", "−0.0005"],
            ["Hourly rate", "−0.2577", "−0.0049"],
        ],
    )
    body(
        doc,
        "For the sampled employee (a stayer), both methods agree on the dominant driver: overtime absence pushes strongly toward staying (SHAP −1.09, LIME −0.125, the largest weight in both). Directional agreement also holds for frequent business travel (−0.43 / −0.07) and the positive pull of income, tenure and age. A repeat LIME run produced near-identical weights, showing acceptable local stability for this sample; residual disagreements on weaker features are expected and discussed as a limitation of local surrogates versus game-theoretic SHAP.",
    )
    body(
        doc,
        "For selected samples, feature-level SHAP values were compared with LIME weights. Perfect numerical equality is not expected—methods differ—but directional agreement on major drivers (for example overtime and compensation-related factors) increases confidence that explanations are not arbitrary. Together, SHAP and LIME address Research Question 2 by improving transparency through dual global and local narratives.",
    )

    add_heading_custom(doc, "4.12 Decision Support Framework", level=2)
    body(
        doc,
        "Prediction and explanation alone do not complete the managerial loop. The proposed framework adds a decision-support layer that translates feature attributions into prioritised, plain-language actions for SME managers.",
    )
    try_add_picture(
        doc,
        "framework_architecture.png",
        width_in=5.9,
        caption="Figure 4.1: Architecture of the XAI managerial decision-support framework (data → preprocessing → models → SHAP/LIME → recommendations → manager UI).",
    )
    body(
        doc,
        "Layers include: (1) data and preprocessing; (2) predictive modelling with multi-model comparison; (3) explainability via SHAP and LIME; (4) decision support—risk score/level, driver lists, recommendation engine and printable managerial report; and (5) delivery through FastAPI and a web dashboard.",
    )
    body(
        doc,
        "The recommendation engine inspects top positive SHAP drivers and raw employee attributes to emit categorised actions spanning workload and burnout, compensation, engagement, flexibility and wellbeing, career progression, financial incentives, and retention maintenance when no critical drivers are present.",
    )
    add_placeholder(doc, "Figure 4.16 Decision support dashboard screenshot (web UI — Risk gauge, SHAP, LIME, Managerial Report)")
    add_caption(
        doc,
        "Figure 4.16: SME manager dashboard showing risk level, explanations and action plan for a selected employee. Paste a screenshot from the running web application.",
    )
    add_para(doc, "Table 4.12: Decision support example — high-risk preset profile", bold=True, size=11, space_after=6)
    add_table(
        doc,
        ["Element", "Content"],
        [
            ["Profile", "Overtime, low pay, low satisfaction, single, no stock options"],
            ["Illustrative attributes", "Age 29; Sales Executive; MonthlyIncome 2,450; OverTime Yes; JobSatisfaction 1"],
            ["Risk band", "High / critical (probability above policy threshold)"],
            ["Example actions", "Cap overtime; compensation review; stay interview; flexibility options"],
        ],
        col_widths=[2.0, 4.5],
    )
    body(
        doc,
        "This closes the loop from black-box score to accountable managerial action, aligning with the Technology Acceptance Model (perceived usefulness and ease of use) and Human–AI Trust Theory (transparency and calibrated trust) developed in Chapter Two.",
    )

    add_heading_custom(doc, "4.13 Discussion of Results", level=2)
    body(
        doc,
        "Objective 1 (AI in SME decisions) is addressed by positioning ML as decision support with explanations that enable judgement. Objective 2 is met by the end-to-end system from raw attributes to recommendations. Objective 3 is met through dual SHAP/LIME explainability. Objective 4 is partially met in an artefact sense—the system operationalises transparency and actionability—while a formal manager trust survey remains future work.",
    )
    body(
        doc,
        "Practically, SMEs should prioritise overtime, compensation and role/level signals; use probability bands plus explanations rather than binary flags only; review false-negative profiles manually; and keep humans in the loop for fairness and context. Limitations include the synthetic IBM dataset, modest recall, absence of a field trust study, and known stability limits of local surrogates such as LIME.",
    )

    add_heading_custom(doc, "4.14 Chapter Summary", level=2)
    body(
        doc,
        "This chapter implemented and evaluated the proposed XAI framework end-to-end. After preprocessing and EDA on the IBM HR attrition dataset, three classifiers were compared; XGBoost offered the best test ROC-AUC (0.804) and precision (0.72). Cross-validation confirmed strong ensemble F1. Feature importance, error analysis, SHAP and LIME provided complementary transparency, and a decision-support layer converted attributions into prioritised managerial actions. Chapter Five summarises contributions, limitations, recommendations and future research.",
    )

    # ===================== CHAPTER 5 =====================
    doc.add_page_break()
    add_heading_custom(doc, "CHAPTER FIVE: CONCLUSION AND RECOMMENDATIONS", level=1)

    add_heading_custom(doc, "5.1 Summary of the Study", level=2)
    body(
        doc,
        "This dissertation designed and validated an Explainable Artificial Intelligence framework that makes machine-learning recommendations more transparent, trustworthy and usable for managerial decision-making in SMEs. Black-box models can deliver useful predictions, yet managers are reluctant to act on outputs they cannot understand or justify. The study combined literature review, pragmatic quantitative experimentation, and full system implementation on the IBM HR attrition case, pairing XGBoost with SHAP and LIME and a recommendation layer for SME managers.",
    )
    body(
        doc,
        "Empirically, XGBoost achieved ROC-AUC 0.8041, accuracy 0.8776, precision 0.7200 and F1 0.5000 on the held-out test set. Explanations highlighted actionable drivers—most notably overtime—while error analysis showed false negatives as the principal operational risk. The prototype dashboard demonstrated how explainability can move from technical plots to managerial reports.",
    )

    add_heading_custom(doc, "5.2 Achievement of Objectives", level=2)
    body(
        doc,
        "Objective 1 was achieved through literature synthesis and conceptual framing of AI’s role under SME constraints. Objective 2 was achieved by building a complete ML pipeline and recommendation engine. Objective 3 was achieved by applying SHAP and LIME for global and local transparency. Objective 4 was partially achieved: the artefact operationalises mechanisms linked to trust and decision quality, but formal field measurement of manager trust was outside the experimental scope and is recommended as future work.",
    )
    add_para(doc, "Table 5.1: Research questions — concise answers", bold=True, size=11, space_after=6)
    add_table(
        doc,
        ["Research question", "Answer based on this study"],
        [
            [
                "RQ1 — Perception/use of AI recommendations",
                "Trust and understandability are prerequisites; the prototype shows a usable risk + explanation + action pattern. Direct manager perception data remain for future work.",
            ],
            [
                "RQ2 — Advantages of SHAP and LIME",
                "SHAP offers consistent global/local attribution; LIME offers accessible case stories; together they reduce black-box opacity.",
            ],
            [
                "RQ3 — Effectiveness for decision quality",
                "Stronger predictive inputs plus explainable actions improve information quality; residual errors require human oversight and threshold policy.",
            ],
        ],
        col_widths=[2.3, 4.2],
    )

    add_heading_custom(doc, "5.3 Contributions", level=2)
    body(
        doc,
        "Theoretical contributions include integrating TAM and Human–AI Trust Theory with dual XAI; shifting evaluation from algorithm-centric accuracy alone toward error costs, local explanations and recommendation quality; and situating explainability in SME managerial constraints. Practical contributions include an end-to-end reference implementation, reusable artefacts, and a decision-support pattern SMEs can adopt without deep ML expertise. Methodological contributions include a leakage-aware pipeline, multi-metric evaluation under imbalance, threshold sensitivity, and dual-explanation audit.",
    )

    add_heading_custom(doc, "5.4 Limitations", level=2)
    body(
        doc,
        "Limitations include: (1) synthetic IBM HR data rather than a live SME census; (2) a single decision domain (attrition); (3) no primary manager survey or field experiment on trust; (4) moderate recall and non-trivial false negatives; (5) explanation limits (LIME stability; SHAP communicative complexity without good UI); (6) library/version dependence of serialised models; and (7) ethics/fairness audits that were acknowledged but not exhaustively tested. These bounds clarify that results are demonstrative and methodological rather than universal organisational proof.",
    )

    add_heading_custom(doc, "5.5 Recommendations", level=2)
    body(
        doc,
        "For SME managers: treat scores as decision inputs not automatic verdicts; require explanations before high-impact actions; align thresholds with business cost of false negatives versus false positives; monitor overtime, compensation equity, role design and career progression; and keep audit logs of model version, score, explanation and human decision.",
    )
    body(
        doc,
        "For developers and consultants: ship predict + explain + recommend together; use train-only resampling; expose probabilities and uncertainty; combine SHAP consistency with LIME or plain-language narratives; and co-design UI language with non-technical users. For policymakers and educators: promote responsible AI literacy in SME digital programmes emphasising explainability and human oversight.",
    )

    add_heading_custom(doc, "5.6 Future Work", level=2)
    body(
        doc,
        "Future work should include field evaluation with SME managers; replication on real organisational data under ethics approval; cost-sensitive and fairness-aware learning; extension to other SME domains; explanation UX studies; MLOps hardening for small IT teams; and counterfactual explanations that answer what would need to change to reduce risk below threshold.",
    )

    add_heading_custom(doc, "5.7 Conclusion", level=2)
    body(
        doc,
        "Black-box AI is a weak foundation for managerial decisions in SMEs, where accountability and scarce specialist capacity make opaque recommendations difficult to trust. This dissertation proposed and implemented an XAI framework that couples competitive predictive performance with SHAP and LIME explainability and a managerial decision-support layer. Using employee attrition as a worked case, XGBoost delivered strong discrimination relative to simpler baselines, explanations surfaced actionable drivers such as overtime, and recommendation logic turned attributions into concrete retention actions.",
    )
    body(
        doc,
        "The lasting contribution is the demonstration that explainable, manager-facing AI is feasible for SMEs when methodology, evaluation and interface design are treated as one system. Future research should take this artefact into live organisational settings to measure the trust and decision-quality gains that theory predicts and that this implementation makes possible.",
    )

    # Short refs note
    add_heading_custom(doc, "Additional references to merge into the main list", level=1)
    body(
        doc,
        "Chen, T. and Guestrin, C. (2016) ‘XGBoost: A scalable tree boosting system’, Proceedings of the 22nd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining, pp. 785–794.",
    )
    body(
        doc,
        "Breiman, L. (2001) ‘Random forests’, Machine Learning, 45(1), pp. 5–32.",
    )
    body(
        doc,
        "Chawla, N.V., Bowyer, K.W., Hall, L.O. and Kegelmeyer, W.P. (2002) ‘SMOTE: Synthetic minority over-sampling technique’, Journal of Artificial Intelligence Research, 16, pp. 321–357.",
    )
    body(
        doc,
        "Pavansubhash (n.d.) IBM HR Analytics Employee Attrition & Performance. Kaggle. Available at: https://www.kaggle.com/datasets/pavansubhasht/ibm-hr-analytics-attrition-dataset",
    )
    body(
        doc,
        "Lundberg, S.M. and Lee, S.I. (2017) and Ribeiro, M.T., Singh, S. and Guestrin, C. (2016) are already in the main dissertation reference list and should remain cited for SHAP and LIME.",
    )

    os.makedirs(os.path.dirname(OUT), exist_ok=True)
    doc.save(OUT)
    print(f"Wrote: {OUT}")
    print(f"Size: {os.path.getsize(OUT)} bytes")


if __name__ == "__main__":
    build()
